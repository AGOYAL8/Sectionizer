from docx import Document
import pandas as pd
from nltk import word_tokenize
from nltk.stem.porter import PorterStemmer
from nltk.stem import WordNetLemmatizer
import re
from nltk.corpus import stopwords
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity
from sklearn.feature_extraction.text import CountVectorizer
from sklearn.feature_extraction.text import HashingVectorizer
from nltk.tokenize.treebank import TreebankWordDetokenizer

import gensim
from gensim.models import Word2Vec

from gensim import utils
from gensim.models.doc2vec import LabeledSentence
from gensim.models.doc2vec import TaggedDocument
from gensim.models import Doc2Vec
from gensim import models
from collections import defaultdict
from gensim import corpora

stemmer = PorterStemmer()
lemmatizer = WordNetLemmatizer()
from gensim import similarities

def read_files(path):
    document = Document(path)
    doc = []
    for para in document.paragraphs:
        if len(para.text) > 101:
            doc.append(para.text)
    df = pd.DataFrame(doc, columns=['sent'])
    return df


def tokenize(x, stem=False, lemma=False):
    x = re.sub(r'(?:<[^>]+>)', '', x)
    pattern = '0-9 $ ` % {} <> -,: _ \ . = +| /'
    x = re.sub(r'(pattern)', '', x)
    tokens = word_tokenize(x)
    tokens = [token for token in tokens if token not in stopwords.words('english')]
    if stem:
        tokens = [stemmer.stem(t) for t in tokens]
    if lemma:
        tokens = [lemmatizer.lemmatize(t) for t in tokens]
    # sent_bigrams = tokens.apply(lambda tweet: [' '.join(bigram) for bigram in ngrams(tweet, 2)])
    return tokens


def preprocess(df, stem=False, lemma=False):
    tokens = []
    if stem:
        tokens = df.sent.apply(lambda x: tokenize(x, True))
    if lemma:
        tokens = df.sent.apply(lambda x: tokenize(x, False, True))
    return tokens

#
def cosine(tokens, stem=False, lemma=False):

    tokens = tokens.apply(lambda x: ' '.join(x))
    for count in range(0, len(tokens)-20):
        sent1 = tokens[count]
        sent2 = tokens[count+1]
        print(count+1, count+2)
        documents = [sent1 + ' ' + sent2]
        # print(documents)
        dictionary = corpora.Dictionary(documents)
        corpus = [dictionary.doc2bow(text) for text in documents]
        lsi = models.LsiModel(corpus, id2word=dictionary, num_topics=2)
        index = similarities.MatrixSimilarity(lsi[corpus])

        index.save('./data/tmp/deerwester.index')
        index = similarities.MatrixSimilarity.load('./data/tmp/deerwester.index')

        sims = index[vec_lsi]  # perform a similarity query against the corpus
        print(list(enumerate(sims)))  # print (document_number, document_similarity) 2-tuples

        # # Create CBOW model
        # model1 = gensim.models.Word2Vec(documents, min_count=1,
        #                                 size=100, window=5)
        # # Create Skip Gram model
        # model2 = gensim.models.Word2Vec(documents, min_count=1, size=100,
        #                                 window=5, sg=1)
        # print(model1.n_similarity(sent1, sent2))
        # print(model2.n_similarity(sent1, sent2))

        # labeled_questions = []
        # labeled_questions.append(TaggedDocument(sent1))
        # labeled_questions.append(TaggedDocument(sent2))

        # model = Doc2Vec(dm=1, min_count=1, window=10, size=150, sample=1e-4, negative=10)
        # model.build_vocab(documents)
        #
        # for epoch in range(20):
        #     model.train(documents, epochs=model.iter, total_examples=model.corpus_count)
        #     print("Epoch #{} is complete.".format(epoch + 1))
        # score = model.n_similarity(sent1, sent2)
        # print(score)


def run(stem=False, lemma=False):
    df_doc = read_files('./data/Report 1_Industry4.0.docx')
    tokens = preprocess(df_doc, stem=stem, lemma=lemma)
    cosine(tokens, stem=stem, lemma=lemma)


def main():
    print('with stemming')
    run(stem=True, lemma=False)
    print('with lemmatizing')
    run(stem=False, lemma=True)


if __name__ == '__main__':
    main()
